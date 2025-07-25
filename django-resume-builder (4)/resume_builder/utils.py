from django.http import HttpResponse
from django.template.loader import render_to_string
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Inches
import io

def generate_pdf_resume(context, template):
    """Generate PDF resume using ReportLab"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    
    personal_info = context.get('personal_info')
    if personal_info:
        story.append(Paragraph(f"{personal_info.first_name} {personal_info.last_name}", title_style))
        contact_info = f"Email: {personal_info.email}"
        if personal_info.phone:
            contact_info += f" | Phone: {personal_info.phone}"
        story.append(Paragraph(contact_info, styles['Normal']))
        story.append(Spacer(1, 12))
        
        if personal_info.summary:
            story.append(Paragraph("Professional Summary", styles['Heading2']))
            story.append(Paragraph(personal_info.summary, styles['Normal']))
            story.append(Spacer(1, 12))
    
    # Work Experience
    work_experiences = context.get('work_experiences', [])
    if work_experiences:
        story.append(Paragraph("Work Experience", styles['Heading2']))
        for exp in work_experiences:
            story.append(Paragraph(f"<b>{exp.position}</b> at {exp.company}", styles['Normal']))
            date_range = f"{exp.start_date}"
            if getattr(exp, 'current', False):
                date_range += " - Present"
            elif exp.end_date:
                date_range += f" - {exp.end_date}"
            story.append(Paragraph(date_range, styles['Normal']))
            story.append(Paragraph(exp.description, styles['Normal']))
            story.append(Spacer(1, 6))
    
    # Education
    educations = context.get('educations', [])
    if educations:
        story.append(Paragraph("Education", styles['Heading2']))
        for edu in educations:
            story.append(Paragraph(f"<b>{edu.degree}</b> in {edu.field_of_study}", styles['Normal']))
            story.append(Paragraph(f"{edu.institution}", styles['Normal']))
            date_range = f"{edu.start_date}"
            if getattr(edu, 'current', False):
                date_range += " - Present"
            elif edu.end_date:
                date_range += f" - {edu.end_date}"
            story.append(Paragraph(date_range, styles['Normal']))
            story.append(Spacer(1, 6))
    
    # Projects
    projects = context.get('projects', [])
    if projects:
        story.append(Paragraph("Projects", styles['Heading2']))
        for project in projects:
            story.append(Paragraph(f"<b>{project.title}</b>", styles['Normal']))
            story.append(Paragraph(project.description, styles['Normal']))
            story.append(Paragraph(f"Technologies: {project.technologies}", styles['Normal']))
            story.append(Spacer(1, 6))
    
    # Technical Skills
    skills = context.get('technical_skills', [])
    if skills:
        story.append(Paragraph("Technical Skills", styles['Heading2']))
        skill_text = ", ".join([f"{skill.name}" for skill in skills])
        story.append(Paragraph(skill_text, styles['Normal']))
    
    try:
        doc.build(story)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="resume_{context["resume"].title}.pdf"'
        return response
    except Exception as e:
        # Fallback to simple text-based PDF
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        
        y = 750
        if personal_info:
            p.setFont("Helvetica-Bold", 16)
            p.drawString(100, y, f"{personal_info.first_name} {personal_info.last_name}")
            y -= 30
            p.setFont("Helvetica", 12)
            p.drawString(100, y, f"Email: {personal_info.email}")
            y -= 20
            if personal_info.phone:
                p.drawString(100, y, f"Phone: {personal_info.phone}")
                y -= 30
        
        p.showPage()
        p.save()
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="resume_{context["resume"].title}.pdf"'
        return response

def generate_docx_resume(context, template):
    """Generate DOCX resume using python-docx"""
    doc = Document()
    
    personal_info = context.get('personal_info')
    if personal_info:
        # Title
        title = doc.add_heading(f"{personal_info.first_name} {personal_info.last_name}", 0)
        title.alignment = 1  # Center alignment
        
        # Contact info
        contact_para = doc.add_paragraph()
        contact_text = f"Email: {personal_info.email}"
        if personal_info.phone:
            contact_text += f" | Phone: {personal_info.phone}"
        contact_para.add_run(contact_text)
        contact_para.alignment = 1
        
        if personal_info.summary:
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(personal_info.summary)
    
    # Work Experience (DOCX)
    work_experiences = context.get('work_experiences', [])
    if work_experiences:
        doc.add_heading('Work Experience', level=1)
        for exp in work_experiences:
            doc.add_paragraph(f"{exp.position} at {exp.company}", style='Heading 2')
            date_range = f"{exp.start_date}"
            if getattr(exp, 'current', False):
                date_range += " - Present"
            elif exp.end_date:
                date_range += f" - {exp.end_date}"
            doc.add_paragraph(date_range)
            doc.add_paragraph(exp.description)
    
    # Education (DOCX)
    educations = context.get('educations', [])
    if educations:
        doc.add_heading('Education', level=1)
        for edu in educations:
            doc.add_paragraph(f"{edu.degree} in {edu.field_of_study}", style='Heading 2')
            doc.add_paragraph(f"{edu.institution}")
            date_range = f"{edu.start_date}"
            if getattr(edu, 'current', False):
                date_range += " - Present"
            elif edu.end_date:
                date_range += f" - {edu.end_date}"
            doc.add_paragraph(date_range)
    
    # Projects
    projects = context.get('projects', [])
    if projects:
        doc.add_heading('Projects', level=1)
        for project in projects:
            doc.add_paragraph(project.title, style='Heading 2')
            doc.add_paragraph(project.description)
            doc.add_paragraph(f"Technologies: {project.technologies}")
    
    # Technical Skills
    skills = context.get('technical_skills', [])
    if skills:
        doc.add_heading('Technical Skills', level=1)
        skill_text = ", ".join([f"{skill.name}" for skill in skills])
        doc.add_paragraph(skill_text)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(), 
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="resume_{context["resume"].title}.docx"'
    return response
